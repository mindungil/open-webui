from asyncio import streams
import hashlib
import json
import logging
import os
import uuid
import re
from datetime import datetime
import html
from functools import lru_cache
from pathlib import Path
from pydub import AudioSegment, effects
from pydub.silence import split_on_silence
from concurrent.futures import ThreadPoolExecutor
from typing import Optional
from transformers import pipeline
import torch
from tqdm import tqdm
import concurrent.futures
import gc

from fnmatch import fnmatch
import aiohttp
import aiofiles
import requests
import mimetypes
from urllib.parse import urljoin, quote

from fastapi import (
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Request,
    UploadFile,
    status,
    APIRouter,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel


from open_webui.utils.auth import get_admin_user, get_verified_user
from open_webui.config import (
    WHISPER_MODEL_AUTO_UPDATE,
    WHISPER_MODEL_DIR,
    CACHE_DIR,
    WHISPER_LANGUAGE,
)

from open_webui.constants import ERROR_MESSAGES
from open_webui.env import (
    AIOHTTP_CLIENT_SESSION_SSL,
    AIOHTTP_CLIENT_TIMEOUT,
    ENV,
    SRC_LOG_LEVELS,
    DEVICE_TYPE,
    ENABLE_FORWARD_USER_INFO_HEADERS,
)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import pipeline as whisper_pipeline


router = APIRouter()

# Constants
MAX_FILE_SIZE_MB = 25
MAX_FILE_SIZE = MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes
AZURE_MAX_FILE_SIZE_MB = 200
AZURE_MAX_FILE_SIZE = AZURE_MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes

log = logging.getLogger(__name__)
log.setLevel(SRC_LOG_LEVELS["AUDIO"])

SPEECH_CACHE_DIR = CACHE_DIR / "audio" / "speech"
SPEECH_CACHE_DIR.mkdir(parents=True, exist_ok=True)


##########################################
#
# Utility functions
#
##########################################

from pydub import AudioSegment
from pydub.utils import mediainfo


def is_audio_conversion_required(file_path):
    """
    Check if the given audio file needs conversion to mp3.
    """
    SUPPORTED_FORMATS = {"flac", "m4a", "mp3", "mp4", "mpeg", "wav", "webm"}

    if not os.path.isfile(file_path):
        log.error(f"File not found: {file_path}")
        return False

    try:
        info = mediainfo(file_path)
        codec_name = info.get("codec_name", "").lower()
        codec_type = info.get("codec_type", "").lower()
        codec_tag_string = info.get("codec_tag_string", "").lower()

        if codec_name == "aac" and codec_type == "audio" and codec_tag_string == "mp4a":
            # File is AAC/mp4a audio, recommend mp3 conversion
            return True

        # If the codec name is in the supported formats
        if codec_name in SUPPORTED_FORMATS:
            return False

        return True
    except Exception as e:
        log.error(f"Error getting audio format: {e}")
        return False


def convert_audio_to_mp3(file_path):
    """Convert audio file to mp3 format."""
    try:
        output_path = os.path.splitext(file_path)[0] + ".mp3"
        audio = AudioSegment.from_file(file_path)
        audio.export(output_path, format="mp3")
        log.info(f"Converted {file_path} to {output_path}")
        return output_path
    except Exception as e:
        log.error(f"Error converting audio file: {e}")
        return None


def set_faster_whisper_model(model: str, auto_update: bool = False):
    whisper_model = None
    if model:
        from faster_whisper import WhisperModel

        # 모델 디렉토리가 존재하는지 확인하고 생성
        if not os.path.exists(WHISPER_MODEL_DIR):
            os.makedirs(WHISPER_MODEL_DIR, exist_ok=True)
            log.info(f"Whisper 모델 디렉토리 생성: {WHISPER_MODEL_DIR}")

        faster_whisper_kwargs = {
            "model_size_or_path": model,
            "device": DEVICE_TYPE if DEVICE_TYPE and DEVICE_TYPE == "cuda" else "cpu",
            "compute_type": "int8",
            "download_root": WHISPER_MODEL_DIR,
            "local_files_only": not auto_update,
        }

        try:
            log.info(f"Faster-Whisper 모델 로딩 시도: {model}")
            whisper_model = WhisperModel(**faster_whisper_kwargs)
            log.info(f"Faster-Whisper 모델 로딩 성공: {model}")
        except Exception as e:
            log.warning(
                f"WhisperModel 초기화 실패 ({model}), local_files_only=False로 재시도: {str(e)}"
            )
            faster_whisper_kwargs["local_files_only"] = False
            try:
                whisper_model = WhisperModel(**faster_whisper_kwargs)
                log.info(f"Faster-Whisper 모델 로딩 성공 (재시도): {model}")
            except Exception as e2:
                log.error(f"Faster-Whisper 모델 로딩 최종 실패: {str(e2)}")
                raise
    return whisper_model


##########################################
#
# Audio API
#
##########################################


class TTSConfigForm(BaseModel):
    OPENAI_API_BASE_URL: str
    OPENAI_API_KEY: str
    OPENAI_PARAMS: Optional[dict] = None
    API_KEY: str
    ENGINE: str
    MODEL: str
    VOICE: str
    SPLIT_ON: str
    AZURE_SPEECH_REGION: str
    AZURE_SPEECH_BASE_URL: str
    AZURE_SPEECH_OUTPUT_FORMAT: str


class STTConfigForm(BaseModel):
    OPENAI_API_BASE_URL: str
    OPENAI_API_KEY: str
    ENGINE: str
    MODEL: str
    SUPPORTED_CONTENT_TYPES: list[str] = []
    WHISPER_MODEL: str
    DEEPGRAM_API_KEY: str
    AZURE_API_KEY: str
    AZURE_REGION: str
    AZURE_LOCALES: str
    AZURE_BASE_URL: str
    AZURE_MAX_SPEAKERS: str


class AudioConfigUpdateForm(BaseModel):
    tts: TTSConfigForm
    stt: STTConfigForm


@router.get("/config")
async def get_audio_config(request: Request, user=Depends(get_admin_user)):
    return {
        "tts": {
            "OPENAI_API_BASE_URL": request.app.state.config.TTS_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.TTS_OPENAI_API_KEY,
            "OPENAI_PARAMS": request.app.state.config.TTS_OPENAI_PARAMS,
            "API_KEY": request.app.state.config.TTS_API_KEY,
            "ENGINE": request.app.state.config.TTS_ENGINE,
            "MODEL": request.app.state.config.TTS_MODEL,
            "VOICE": request.app.state.config.TTS_VOICE,
            "SPLIT_ON": request.app.state.config.TTS_SPLIT_ON,
            "AZURE_SPEECH_REGION": request.app.state.config.TTS_AZURE_SPEECH_REGION,
            "AZURE_SPEECH_BASE_URL": request.app.state.config.TTS_AZURE_SPEECH_BASE_URL,
            "AZURE_SPEECH_OUTPUT_FORMAT": request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT,
        },
        "stt": {
            "OPENAI_API_BASE_URL": request.app.state.config.STT_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.STT_OPENAI_API_KEY,
            "ENGINE": request.app.state.config.STT_ENGINE,
            "MODEL": request.app.state.config.STT_MODEL,
            "SUPPORTED_CONTENT_TYPES": request.app.state.config.STT_SUPPORTED_CONTENT_TYPES,
            "WHISPER_MODEL": request.app.state.config.WHISPER_MODEL,
            "DEEPGRAM_API_KEY": request.app.state.config.DEEPGRAM_API_KEY,
            "AZURE_API_KEY": request.app.state.config.AUDIO_STT_AZURE_API_KEY,
            "AZURE_REGION": request.app.state.config.AUDIO_STT_AZURE_REGION,
            "AZURE_LOCALES": request.app.state.config.AUDIO_STT_AZURE_LOCALES,
            "AZURE_BASE_URL": request.app.state.config.AUDIO_STT_AZURE_BASE_URL,
            "AZURE_MAX_SPEAKERS": request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS,
        },
    }


@router.post("/config/update")
async def update_audio_config(
    request: Request, form_data: AudioConfigUpdateForm, user=Depends(get_admin_user)
):
    request.app.state.config.TTS_OPENAI_API_BASE_URL = form_data.tts.OPENAI_API_BASE_URL
    request.app.state.config.TTS_OPENAI_API_KEY = form_data.tts.OPENAI_API_KEY
    request.app.state.config.TTS_OPENAI_PARAMS = form_data.tts.OPENAI_PARAMS
    request.app.state.config.TTS_API_KEY = form_data.tts.API_KEY
    request.app.state.config.TTS_ENGINE = form_data.tts.ENGINE
    request.app.state.config.TTS_MODEL = form_data.tts.MODEL
    request.app.state.config.TTS_VOICE = form_data.tts.VOICE
    request.app.state.config.TTS_SPLIT_ON = form_data.tts.SPLIT_ON
    request.app.state.config.TTS_AZURE_SPEECH_REGION = form_data.tts.AZURE_SPEECH_REGION
    request.app.state.config.TTS_AZURE_SPEECH_BASE_URL = (
        form_data.tts.AZURE_SPEECH_BASE_URL
    )
    request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT = (
        form_data.tts.AZURE_SPEECH_OUTPUT_FORMAT
    )

    request.app.state.config.STT_OPENAI_API_BASE_URL = form_data.stt.OPENAI_API_BASE_URL
    request.app.state.config.STT_OPENAI_API_KEY = form_data.stt.OPENAI_API_KEY
    request.app.state.config.STT_ENGINE = form_data.stt.ENGINE
    request.app.state.config.STT_MODEL = form_data.stt.MODEL
    request.app.state.config.STT_SUPPORTED_CONTENT_TYPES = (
        form_data.stt.SUPPORTED_CONTENT_TYPES
    )

    request.app.state.config.WHISPER_MODEL = form_data.stt.WHISPER_MODEL
    request.app.state.config.DEEPGRAM_API_KEY = form_data.stt.DEEPGRAM_API_KEY
    request.app.state.config.AUDIO_STT_AZURE_API_KEY = form_data.stt.AZURE_API_KEY
    request.app.state.config.AUDIO_STT_AZURE_REGION = form_data.stt.AZURE_REGION
    request.app.state.config.AUDIO_STT_AZURE_LOCALES = form_data.stt.AZURE_LOCALES
    request.app.state.config.AUDIO_STT_AZURE_BASE_URL = form_data.stt.AZURE_BASE_URL
    request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS = (
        form_data.stt.AZURE_MAX_SPEAKERS
    )

    if request.app.state.config.STT_ENGINE == "":
        request.app.state.faster_whisper_model = set_faster_whisper_model(
            form_data.stt.WHISPER_MODEL, WHISPER_MODEL_AUTO_UPDATE
        )
    else:
        request.app.state.faster_whisper_model = None

    return {
        "tts": {
            "ENGINE": request.app.state.config.TTS_ENGINE,
            "MODEL": request.app.state.config.TTS_MODEL,
            "VOICE": request.app.state.config.TTS_VOICE,
            "OPENAI_API_BASE_URL": request.app.state.config.TTS_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.TTS_OPENAI_API_KEY,
            "OPENAI_PARAMS": request.app.state.config.TTS_OPENAI_PARAMS,
            "API_KEY": request.app.state.config.TTS_API_KEY,
            "SPLIT_ON": request.app.state.config.TTS_SPLIT_ON,
            "AZURE_SPEECH_REGION": request.app.state.config.TTS_AZURE_SPEECH_REGION,
            "AZURE_SPEECH_BASE_URL": request.app.state.config.TTS_AZURE_SPEECH_BASE_URL,
            "AZURE_SPEECH_OUTPUT_FORMAT": request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT,
        },
        "stt": {
            "OPENAI_API_BASE_URL": request.app.state.config.STT_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.STT_OPENAI_API_KEY,
            "ENGINE": request.app.state.config.STT_ENGINE,
            "MODEL": request.app.state.config.STT_MODEL,
            "SUPPORTED_CONTENT_TYPES": request.app.state.config.STT_SUPPORTED_CONTENT_TYPES,
            "WHISPER_MODEL": request.app.state.config.WHISPER_MODEL,
            "DEEPGRAM_API_KEY": request.app.state.config.DEEPGRAM_API_KEY,
            "AZURE_API_KEY": request.app.state.config.AUDIO_STT_AZURE_API_KEY,
            "AZURE_REGION": request.app.state.config.AUDIO_STT_AZURE_REGION,
            "AZURE_LOCALES": request.app.state.config.AUDIO_STT_AZURE_LOCALES,
            "AZURE_BASE_URL": request.app.state.config.AUDIO_STT_AZURE_BASE_URL,
            "AZURE_MAX_SPEAKERS": request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS,
        },
    }


def load_speech_pipeline(request):
    from transformers import pipeline
    from datasets import load_dataset

    if request.app.state.speech_synthesiser is None:
        request.app.state.speech_synthesiser = pipeline(
            "text-to-speech", "microsoft/speecht5_tts"
        )

    if request.app.state.speech_speaker_embeddings_dataset is None:
        request.app.state.speech_speaker_embeddings_dataset = load_dataset(
            "Matthijs/cmu-arctic-xvectors", split="validation"
        )


@router.post("/speech")
async def speech(request: Request, user=Depends(get_verified_user)):
    body = await request.body()
    name = hashlib.sha256(
        body
        + str(request.app.state.config.TTS_ENGINE).encode("utf-8")
        + str(request.app.state.config.TTS_MODEL).encode("utf-8")
    ).hexdigest()

    file_path = SPEECH_CACHE_DIR.joinpath(f"{name}.mp3")
    file_body_path = SPEECH_CACHE_DIR.joinpath(f"{name}.json")

    # Check if the file already exists in the cache
    if file_path.is_file():
        return FileResponse(file_path)

    payload = None
    try:
        payload = json.loads(body.decode("utf-8"))
    except Exception as e:
        log.exception(e)
        raise HTTPException(status_code=400, detail="Invalid JSON payload")

    r = None
    if request.app.state.config.TTS_ENGINE == "openai":
        payload["model"] = request.app.state.config.TTS_MODEL

        try:
            timeout = aiohttp.ClientTimeout(total=AIOHTTP_CLIENT_TIMEOUT)
            async with aiohttp.ClientSession(
                timeout=timeout, trust_env=True
            ) as session:
                payload = {
                    **payload,
                    **(request.app.state.config.TTS_OPENAI_PARAMS or {}),
                }

                r = await session.post(
                    url=f"{request.app.state.config.TTS_OPENAI_API_BASE_URL}/audio/speech",
                    json=payload,
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {request.app.state.config.TTS_OPENAI_API_KEY}",
                        **(
                            {
                                "X-OpenWebUI-User-Name": quote(user.name, safe=" "),
                                "X-OpenWebUI-User-Id": user.id,
                                "X-OpenWebUI-User-Email": user.email,
                                "X-OpenWebUI-User-Role": user.role,
                            }
                            if ENABLE_FORWARD_USER_INFO_HEADERS
                            else {}
                        ),
                    },
                    ssl=AIOHTTP_CLIENT_SESSION_SSL,
                )

                r.raise_for_status()

                async with aiofiles.open(file_path, "wb") as f:
                    await f.write(await r.read())

                async with aiofiles.open(file_body_path, "w") as f:
                    await f.write(json.dumps(payload))

            return FileResponse(file_path)

        except Exception as e:
            log.exception(e)
            detail = None

            status_code = 500
            detail = f"Open WebUI: Server Connection Error"

            if r is not None:
                status_code = r.status

                try:
                    res = await r.json()
                    if "error" in res:
                        detail = f"External: {res['error']}"
                except Exception:
                    detail = f"External: {e}"

            raise HTTPException(
                status_code=status_code,
                detail=detail,
            )

    elif request.app.state.config.TTS_ENGINE == "elevenlabs":
        voice_id = payload.get("voice", "")

        if voice_id not in get_available_voices(request):
            raise HTTPException(
                status_code=400,
                detail="Invalid voice id",
            )

        try:
            timeout = aiohttp.ClientTimeout(total=AIOHTTP_CLIENT_TIMEOUT)
            async with aiohttp.ClientSession(
                timeout=timeout, trust_env=True
            ) as session:
                async with session.post(
                    f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}",
                    json={
                        "text": payload["input"],
                        "model_id": request.app.state.config.TTS_MODEL,
                        "voice_settings": {"stability": 0.5, "similarity_boost": 0.5},
                    },
                    headers={
                        "Accept": "audio/mpeg",
                        "Content-Type": "application/json",
                        "xi-api-key": request.app.state.config.TTS_API_KEY,
                    },
                    ssl=AIOHTTP_CLIENT_SESSION_SSL,
                ) as r:
                    r.raise_for_status()

                    async with aiofiles.open(file_path, "wb") as f:
                        await f.write(await r.read())

                    async with aiofiles.open(file_body_path, "w") as f:
                        await f.write(json.dumps(payload))

            return FileResponse(file_path)

        except Exception as e:
            log.exception(e)
            detail = None

            try:
                if r.status != 200:
                    res = await r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
            except Exception:
                detail = f"External: {e}"

            raise HTTPException(
                status_code=getattr(r, "status", 500) if r else 500,
                detail=detail if detail else "Open WebUI: Server Connection Error",
            )

    elif request.app.state.config.TTS_ENGINE == "azure":
        try:
            payload = json.loads(body.decode("utf-8"))
        except Exception as e:
            log.exception(e)
            raise HTTPException(status_code=400, detail="Invalid JSON payload")

        region = request.app.state.config.TTS_AZURE_SPEECH_REGION or "eastus"
        base_url = request.app.state.config.TTS_AZURE_SPEECH_BASE_URL
        language = request.app.state.config.TTS_VOICE
        locale = "-".join(request.app.state.config.TTS_VOICE.split("-")[:1])
        output_format = request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT

        try:
            data = f"""<speak version="1.0" xmlns="http://www.w3.org/2001/10/synthesis" xml:lang="{locale}">
                <voice name="{language}">{html.escape(payload["input"])}</voice>
            </speak>"""
            timeout = aiohttp.ClientTimeout(total=AIOHTTP_CLIENT_TIMEOUT)
            async with aiohttp.ClientSession(
                timeout=timeout, trust_env=True
            ) as session:
                async with session.post(
                    (base_url or f"https://{region}.tts.speech.microsoft.com")
                    + "/cognitiveservices/v1",
                    headers={
                        "Ocp-Apim-Subscription-Key": request.app.state.config.TTS_API_KEY,
                        "Content-Type": "application/ssml+xml",
                        "X-Microsoft-OutputFormat": output_format,
                    },
                    data=data,
                    ssl=AIOHTTP_CLIENT_SESSION_SSL,
                ) as r:
                    r.raise_for_status()

                    async with aiofiles.open(file_path, "wb") as f:
                        await f.write(await r.read())

                    async with aiofiles.open(file_body_path, "w") as f:
                        await f.write(json.dumps(payload))

                    return FileResponse(file_path)

        except Exception as e:
            log.exception(e)
            detail = None

            try:
                if r.status != 200:
                    res = await r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
            except Exception:
                detail = f"External: {e}"

            raise HTTPException(
                status_code=getattr(r, "status", 500) if r else 500,
                detail=detail if detail else "Open WebUI: Server Connection Error",
            )

    elif request.app.state.config.TTS_ENGINE == "transformers":
        payload = None
        try:
            payload = json.loads(body.decode("utf-8"))
        except Exception as e:
            log.exception(e)
            raise HTTPException(status_code=400, detail="Invalid JSON payload")

        import torch
        import soundfile as sf

        load_speech_pipeline(request)

        embeddings_dataset = request.app.state.speech_speaker_embeddings_dataset

        speaker_index = 6799
        try:
            speaker_index = embeddings_dataset["filename"].index(
                request.app.state.config.TTS_MODEL
            )
        except Exception:
            pass

        speaker_embedding = torch.tensor(
            embeddings_dataset[speaker_index]["xvector"]
        ).unsqueeze(0)

        speech = request.app.state.speech_synthesiser(
            payload["input"],
            forward_params={"speaker_embeddings": speaker_embedding},
        )

        sf.write(file_path, speech["audio"], samplerate=speech["sampling_rate"])

        async with aiofiles.open(file_body_path, "w") as f:
            await f.write(json.dumps(payload))

        return FileResponse(file_path)

def convert_to_wav(input_path):
        try:
            audio = AudioSegment.from_file(input_path)
            audio = effects.normalize(audio)
            output_path = os.path.splitext(input_path)[0] + ".wav"
            audio.export(output_path, format="wav")
            log.info(f"WAV 변환 완료: {output_path}")
            return output_path
        except Exception as e:
            log.error(f"WAV 변환 실패: {str(e)}")
            raise

# 무음을 기준으로 오디오 분할 (짧은 청크 병합, 긴 청크 슬라이스)
# --> 적응형 라이브러리 있어서 변경 예정
def split_audio(audio_path, min_silence_len=500, silence_thresh=-40, max_chunk_ms=30000, min_chunk_ms=3000):
    try:
        audio = AudioSegment.from_wav(audio_path)

        # 무음 기준 분할
        raw_chunks = split_on_silence(
            audio,
            min_silence_len=min_silence_len,
            silence_thresh=silence_thresh
        )

        # 너무 짧은 청크 합치기
        def merge_short_chunks(chunks):
            merged = []
            buffer = None
            for chunk in chunks:
                if len(chunk) < min_chunk_ms:
                    if buffer is not None:
                        buffer += chunk
                    else:
                        buffer = chunk
                else:
                    if buffer is not None:
                        chunk = buffer + chunk
                        buffer = None
                    merged.append(chunk)
            if buffer is not None:
                merged.append(buffer)
            return merged

        merged_chunks = merge_short_chunks(raw_chunks)

        # 너무 긴 청크는 나누기
        final_chunks = []
        for chunk in merged_chunks:
            if len(chunk) > max_chunk_ms:
                for i in range(0, len(chunk), max_chunk_ms):
                    sliced = chunk[i:i + max_chunk_ms]
                    if len(sliced) >= min_chunk_ms:
                        final_chunks.append(sliced)
            else:
                final_chunks.append(chunk)

        # 분할된 청크 WAV로 저장
        chunk_paths = []
        for i, chunk in enumerate(final_chunks):
            chunk_path = f"{os.path.splitext(audio_path)[0]}_chunk_sil_{i}.wav"
            chunk.export(chunk_path, format="wav")
            chunk_paths.append(chunk_path)

        log.info(f"총 {len(chunk_paths)}개의 청크로 분할 완료 (무음 기준 + 병합 + 슬라이스)")
        return chunk_paths

    except Exception as e:
        log.error(f"오디오 분할 실패: {str(e)}")
        raise

# 한 청크 Whisper STT 수행
def process_chunk(chunk_path, pipe):
    try:
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gc.collect()
        result = pipe(chunk_path)
        return result.get('text', '').strip()
    except Exception as e:
        log.error(f"청크 처리 실패 ({chunk_path}): {str(e)}")
        return ""

def convert_seconds_to_hms(seconds):
    """초를 HH:MM:SS 형식으로 변환 (분 단위로 반올림)"""
    # 회의록에서는 초 단위까지 필요없으므로 분 단위로 반올림
    total_seconds = round(seconds)
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60

    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    else:
        return f"{minutes:02d}:{seconds:02d}"

# 전체 진행 함수
def transcribe_long_audio(request: Request, file_path, model_name='large-v3'):
    try:
        # GPU 리소스 정리
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gc.collect()
            torch.backends.cudnn.benchmark = True

        # 오디오 WAV로 변환
        wav_path = convert_to_wav(file_path)

        # 디바이스 설정 (GPU or CPU)
        device = 0 if torch.cuda.is_available() else -1
        log.info(f"Faster-Whisper 실행 중 (Device: {'GPU' if device == 0 else 'CPU'})")

        # Faster-Whisper 모델 로딩
        log.info("Faster-Whisper 모델 로딩 중...")

        # 설정에서 모델 가져오기 (기본값: large-v3 모델)
        whisper_model_name = request.app.state.config.WHISPER_MODEL or "large-v3"
        log.info(f"사용할 모델: {whisper_model_name}")

        # 기존 모델이 있으면 사용, 없으면 새로 생성
        if request.app.state.faster_whisper_model is None:
            try:
                request.app.state.faster_whisper_model = set_faster_whisper_model(
                    whisper_model_name,
                    WHISPER_MODEL_AUTO_UPDATE
                )
            except Exception as e:
                log.warning(f"설정된 모델 로딩 실패, 기본 모델로 시도: {str(e)}")
                # 기본 모델로 재시도
                request.app.state.faster_whisper_model = set_faster_whisper_model(
                    "large-v3",
                    True  # 자동 업데이트 활성화
                )

        whisper = request.app.state.faster_whisper_model

        if whisper is None:
            raise RuntimeError("Faster-Whisper 모델을 로드할 수 없습니다.")

        # 회의록에 최적화된 VAD 설정으로 전체 파일 처리
        log.info("회의록 최적화 설정으로 Faster-Whisper 처리 중...")
        segments, info = whisper.transcribe(
            wav_path,
            beam_size=5,
            vad_filter=True,  # VAD 활성화
            vad_parameters=dict(
                min_silence_duration_ms=1000,  # 5초 무음 (기존 3초→5초, 더 엄격하게)
                speech_pad_ms=200,             # 0.2초 패딩 (기존 0.5초→0.2초, 더 정확하게)
                threshold=0.7                  # 0.7 임계값 (기존 0.5→0.7, 더 엄격하게)
            ),
            language="ko",
            # 회의록 품질 향상을 위한 추가 설정
            condition_on_previous_text=False,  # 이전 텍스트 의존성 제거
            compression_ratio_threshold=2.4,   # 반복 패턴 감지
            no_speech_threshold=0.7,          # 무음 구간 더 엄격하게 (기존 0.6→0.7)
            temperature=0.0,                  # 일관성 있는 출력
            initial_prompt="이것은 한국어 회의 녹음입니다. 정확하고 자연스러운 문장으로 전사해주세요."
        )

        log.info(f"언어 감지: {info.language} (확률: {info.language_probability:.2f})")

        # 세그먼트 후처리 및 병합
        merged_segments = merge_short_segments(segments)

        # 결과 처리
        results = []
        segments_list = []

        for segment in merged_segments:
            text = segment['text']
            if text and len(text) > 3:  # 3글자 이상만 포함
                # 반복 패턴 제거
                text = filter_repetitive_text(text)
                if text:  # 필터링 후에도 텍스트가 남아있으면
                    results.append(text)
                    segments_list.append({
                        "start": convert_seconds_to_hms(segment['start']),
                        "end": convert_seconds_to_hms(segment['end']),
                        "text": text
                    })

        # 전체 텍스트 통합
        plain_text = " ".join(results)

        # 향상된 docx 문서 생성
        doc = create_enhanced_meeting_document(segments_list)

        log.info("Faster-Whisper STT 처리 완료")
        return {
            'plain_text': plain_text,
            'docx_document': doc,
            'segments': segments_list
        }

    except Exception as e:
        log.error(f"Faster-Whisper STT 처리 실패: {str(e)}")
        raise

def transcription_handler(request, file_path, metadata):
    filename = os.path.basename(file_path)
    file_dir = os.path.dirname(file_path)
    id = filename.split(".")[0]

    metadata = metadata or {}

    languages = [
        metadata.get("language", None) if not WHISPER_LANGUAGE else WHISPER_LANGUAGE,
        None,  # Always fallback to None in case transcription fails
    ]

    if request.app.state.config.STT_ENGINE == "":
        if request.app.state.faster_whisper_model is None:
            request.app.state.faster_whisper_model = set_faster_whisper_model(
                request.app.state.config.WHISPER_MODEL
            )

        model = request.app.state.faster_whisper_model
        segments, info = model.transcribe(
            file_path,
            beam_size=5,
            vad_filter=request.app.state.config.WHISPER_VAD_FILTER,
            language=languages[0],
        )
        log.info(
            "Detected language '%s' with probability %f"
            % (info.language, info.language_probability)
        )

        # 시간대별 세그먼트 정보를 포함한 데이터 생성
        segments_list = []
        for segment in segments:
            segments_list.append({
                "start": round(segment.start, 2),
                "end": round(segment.end, 2),
                "text": segment.text.strip()
            })

        data = {
            "text": "".join([segment["text"] for segment in segments_list]),
            "segments": segments_list
        }

        # save the transcript to a json file
        transcript_file = f"{file_dir}/{id}.json"
        with open(transcript_file, "w") as f:
            json.dump(data, f)

        log.debug(data)
        return data

    elif request.app.state.config.STT_ENGINE == "openai":
        r = None
        try:
            for language in languages:
                payload = {
                    "model": request.app.state.config.STT_MODEL,
                }

                if language:
                    payload["language"] = language

                r = requests.post(
                    url=f"{request.app.state.config.STT_OPENAI_API_BASE_URL}/audio/transcriptions",
                    headers={
                        "Authorization": f"Bearer {request.app.state.config.STT_OPENAI_API_KEY}"
                    },
                    files={"file": (filename, open(file_path, "rb"))},
                    data=payload,
                )

                if r.status_code == 200:
                    # Successful transcription
                    break

            r.raise_for_status()
            data = r.json()

            # save the transcript to a json file
            transcript_file = f"{file_dir}/{id}.json"
            with open(transcript_file, "w") as f:
                json.dump(data, f)

            return data
        except Exception as e:
            log.exception(e)

            detail = None
            if r is not None:
                try:
                    res = r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
                except Exception:
                    detail = f"External: {e}"

            raise Exception(detail if detail else "Open WebUI: Server Connection Error")

    elif request.app.state.config.STT_ENGINE == "deepgram":
        try:
            # Determine the MIME type of the file
            mime, _ = mimetypes.guess_type(file_path)
            if not mime:
                mime = "audio/wav"  # fallback to wav if undetectable

            # Read the audio file
            with open(file_path, "rb") as f:
                file_data = f.read()

            # Build headers and parameters
            headers = {
                "Authorization": f"Token {request.app.state.config.DEEPGRAM_API_KEY}",
                "Content-Type": mime,
            }

            for language in languages:
                params = {}
                if request.app.state.config.STT_MODEL:
                    params["model"] = request.app.state.config.STT_MODEL

                if language:
                    params["language"] = language

                # Make request to Deepgram API
                r = requests.post(
                    "https://api.deepgram.com/v1/listen?smart_format=true",
                    headers=headers,
                    params=params,
                    data=file_data,
                )

                if r.status_code == 200:
                    # Successful transcription
                    break

            r.raise_for_status()
            response_data = r.json()

            # Extract transcript from Deepgram response
            try:
                transcript = response_data["results"]["channels"][0]["alternatives"][
                    0
                ].get("transcript", "")
            except (KeyError, IndexError) as e:
                log.error(f"Malformed response from Deepgram: {str(e)}")
                raise Exception(
                    "Failed to parse Deepgram response - unexpected response format"
                )
            data = {"text": transcript.strip()}

            # Save transcript
            transcript_file = f"{file_dir}/{id}.json"
            with open(transcript_file, "w") as f:
                json.dump(data, f)

            return data

        except Exception as e:
            log.exception(e)
            detail = None
            if r is not None:
                try:
                    res = r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
                except Exception:
                    detail = f"External: {e}"
            raise Exception(detail if detail else "Open WebUI: Server Connection Error")

    elif request.app.state.config.STT_ENGINE == "azure":
        # Check file exists and size
        if not os.path.exists(file_path):
            raise HTTPException(status_code=400, detail="Audio file not found")

        # Check file size (Azure has a larger limit of 200MB)
        file_size = os.path.getsize(file_path)
        if file_size > AZURE_MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds Azure's limit of {AZURE_MAX_FILE_SIZE_MB}MB",
            )

        api_key = request.app.state.config.AUDIO_STT_AZURE_API_KEY
        region = request.app.state.config.AUDIO_STT_AZURE_REGION or "eastus"
        locales = request.app.state.config.AUDIO_STT_AZURE_LOCALES
        base_url = request.app.state.config.AUDIO_STT_AZURE_BASE_URL
        max_speakers = request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS or 3

        # IF NO LOCALES, USE DEFAULTS
        if len(locales) < 2:
            locales = [
                "en-US",
                "es-ES",
                "es-MX",
                "fr-FR",
                "hi-IN",
                "it-IT",
                "de-DE",
                "en-GB",
                "en-IN",
                "ja-JP",
                "ko-KR",
                "pt-BR",
                "zh-CN",
            ]
            locales = ",".join(locales)

        if not api_key or not region:
            raise HTTPException(
                status_code=400,
                detail="Azure API key is required for Azure STT",
            )

        r = None
        try:
            # Prepare the request
            data = {
                "definition": json.dumps(
                    {
                        "locales": locales.split(","),
                        "diarization": {"maxSpeakers": max_speakers, "enabled": True},
                    }
                    if locales
                    else {}
                )
            }

            url = (
                base_url or f"https://{region}.api.cognitive.microsoft.com"
            ) + "/speechtotext/transcriptions:transcribe?api-version=2024-11-15"

            # Use context manager to ensure file is properly closed
            with open(file_path, "rb") as audio_file:
                r = requests.post(
                    url=url,
                    files={"audio": audio_file},
                    data=data,
                    headers={
                        "Ocp-Apim-Subscription-Key": api_key,
                    },
                )

            r.raise_for_status()
            response = r.json()

            # Extract transcript from response
            if not response.get("combinedPhrases"):
                raise ValueError("No transcription found in response")

            # Get the full transcript from combinedPhrases
            transcript = response["combinedPhrases"][0].get("text", "").strip()
            if not transcript:
                raise ValueError("Empty transcript in response")

            data = {"text": transcript}

            # Save transcript to json file (consistent with other providers)
            transcript_file = f"{file_dir}/{id}.json"
            with open(transcript_file, "w") as f:
                json.dump(data, f)

            log.debug(data)
            return data

        except (KeyError, IndexError, ValueError) as e:
            log.exception("Error parsing Azure response")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse Azure response: {str(e)}",
            )
        except requests.exceptions.RequestException as e:
            log.exception(e)
            detail = None

            try:
                if r is not None and r.status_code != 200:
                    res = r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
            except Exception:
                detail = f"External: {e}"

            raise HTTPException(
                status_code=getattr(r, "status_code", 500) if r else 500,
                detail=detail if detail else "Open WebUI: Server Connection Error",
            )

def transcribe(request: Request, file_path: str, metadata: Optional[dict] = None, filedata: list = None):
    log.info(f"transcribe: {file_path} {metadata}")
    log.info(f"filedata: {filedata}")

    result = transcribe_long_audio(request, file_path)
    plain_text = result['plain_text']
    docx_doc = result['docx_document']
    segments = result['segments']

    try:
        # 1. 기존 txt 파일 저장 (변경 없음)
        save_path = os.path.join(os.path.dirname(file_path), f"{os.path.splitext(filedata[1])[0]}.txt")
        with open(save_path, "w", encoding="utf-8") as f:
            f.write(plain_text)
        log.info(f"Transcript saved to: {save_path}")

        from open_webui.models.files import FileForm, Files
        from open_webui.storage.provider import Storage
        import uuid

        # 2. txt 파일을 영구 저장소에 업로드 (변경 없음)
        txt_id = f"{filedata[0]}txt"
        txt_name = f"{os.path.splitext(filedata[1])[0]}txt"
        txt_filename = f"{filedata[2]}txt"
        with open(save_path, "rb") as f:
            txt_bytes, txt_storage_path = Storage.upload_file(
                f, txt_filename,
                tags = {
                    **filedata[3],
                    "OpenWebUI-File-Id": txt_id,
                    "OpenWebUI-Transcript-Of": f"{filedata[0]}"
                }
            )

        user_id = filedata[3]["OpenWebUI-User-Id"]
        # 3. Files 테이블에 txt 파일 row 생성 (변경 없음)
        Files.insert_new_file(
            user_id,
            FileForm(
                id = txt_id,
                filename = txt_name,
                path = txt_storage_path,
                meta = {
                    "name": txt_name,
                    "content_type": "text/plain",
                    "size": len(txt_bytes),
                    "data": {"transcript_of": filedata[0]}
                },
            )
        )

        # 4. docx 파일 저장 및 업로드 (새로 추가)
        docx_path = os.path.join(os.path.dirname(file_path), f"{os.path.splitext(filedata[1])[0]}.docx")
        docx_doc.save(docx_path)
        log.info(f"Detailed transcript saved to: {docx_path}")

        # 5. docx 파일을 영구 저장소에 업로드
        docx_id = f"{filedata[0]}docx"
        docx_name = f"{os.path.splitext(filedata[1])[0]}.docx"
        docx_filename = f"{filedata[2]}.docx"
        with open(docx_path, "rb") as f:
            docx_bytes, docx_storage_path = Storage.upload_file(
                f, docx_filename,
                tags = {
                    **filedata[3],
                    "OpenWebUI-File-Id": docx_id,
                    "OpenWebUI-Transcript-Of": f"{filedata[0]}",
                    "OpenWebUI-Transcript-Type": "detailed"
                }
            )

        # 6. Files 테이블에 docx 파일 row 생성
        Files.insert_new_file(
            user_id,
            FileForm(
                id = docx_id,
                filename = docx_name,
                path = docx_storage_path,
                meta = {
                    "name": docx_name,
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "size": len(docx_bytes),
                    "data": {
                        "transcript_of": filedata[0],
                        "transcript_type": "detailed",
                        "segments": segments
                    }
                },
            )
        )

        return {
            "text": plain_text,
            "transcript_file_id": txt_id,
            "detailed_transcript_file_id": docx_id
        }

    except Exception as e:
        log.exception(f"Failed to save or upload transcript: {e}")
        return {
            "text": plain_text,
            "error": f"Transcript saved locally, but upload failed: {str(e)}",
        }



@router.post("/transcriptions")
def transcription(
    request: Request,
    file: UploadFile = File(...),
    language: Optional[str] = Form(None),
    user=Depends(get_verified_user),
):
    log.info(f"file.content_type: {file.content_type}")

    stt_supported_content_types = getattr(
        request.app.state.config, "STT_SUPPORTED_CONTENT_TYPES", []
    )

    if not any(
        fnmatch(file.content_type, content_type)
        for content_type in (
            stt_supported_content_types
            if stt_supported_content_types
            and any(t.strip() for t in stt_supported_content_types)
            else ["audio/*", "video/webm"]
        )
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.FILE_NOT_SUPPORTED,
        )

    try:
        ext = file.filename.split(".")[-1]
        id = uuid.uuid4()

        filename = f"{id}.{ext}"
        contents = file.file.read()

        file_dir = f"{CACHE_DIR}/audio/transcriptions"
        os.makedirs(file_dir, exist_ok=True)
        file_path = f"{file_dir}/{filename}"

        with open(file_path, "wb") as f:
            f.write(contents)

        try:
            metadata = None

            if language:
                metadata = {"language": language}

            result = transcribe(request, file_path, metadata)

            return {
                **result,
                "filename": os.path.basename(file_path),
            }

        except Exception as e:
            log.exception(e)

            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT(e),
            )

    except Exception as e:
        log.exception(e)

        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.DEFAULT(e),
        )


def get_available_models(request: Request) -> list[dict]:
    available_models = []
    if request.app.state.config.TTS_ENGINE == "openai":
        # Use custom endpoint if not using the official OpenAI API URL
        if not request.app.state.config.TTS_OPENAI_API_BASE_URL.startswith(
            "https://api.openai.com"
        ):
            try:
                response = requests.get(
                    f"{request.app.state.config.TTS_OPENAI_API_BASE_URL}/audio/models"
                )
                response.raise_for_status()
                data = response.json()
                available_models = data.get("models", [])
            except Exception as e:
                log.error(f"Error fetching models from custom endpoint: {str(e)}")
                available_models = [{"id": "tts-1"}, {"id": "tts-1-hd"}]
        else:
            available_models = [{"id": "tts-1"}, {"id": "tts-1-hd"}]
    elif request.app.state.config.TTS_ENGINE == "elevenlabs":
        try:
            response = requests.get(
                "https://api.elevenlabs.io/v1/models",
                headers={
                    "xi-api-key": request.app.state.config.TTS_API_KEY,
                    "Content-Type": "application/json",
                },
                timeout=5,
            )
            response.raise_for_status()
            models = response.json()

            available_models = [
                {"name": model["name"], "id": model["model_id"]} for model in models
            ]
        except requests.RequestException as e:
            log.error(f"Error fetching voices: {str(e)}")
    return available_models


@router.get("/models")
async def get_models(request: Request, user=Depends(get_verified_user)):
    return {"models": get_available_models(request)}


def get_available_voices(request) -> dict:
    """Returns {voice_id: voice_name} dict"""
    available_voices = {}
    if request.app.state.config.TTS_ENGINE == "openai":
        # Use custom endpoint if not using the official OpenAI API URL
        if not request.app.state.config.TTS_OPENAI_API_BASE_URL.startswith(
            "https://api.openai.com"
        ):
            try:
                response = requests.get(
                    f"{request.app.state.config.TTS_OPENAI_API_BASE_URL}/audio/voices"
                )
                response.raise_for_status()
                data = response.json()
                voices_list = data.get("voices", [])
                available_voices = {voice["id"]: voice["name"] for voice in voices_list}
            except Exception as e:
                log.error(f"Error fetching voices from custom endpoint: {str(e)}")
                available_voices = {
                    "alloy": "alloy",
                    "echo": "echo",
                    "fable": "fable",
                    "onyx": "onyx",
                    "nova": "nova",
                    "shimmer": "shimmer",
                }
        else:
            available_voices = {
                "alloy": "alloy",
                "echo": "echo",
                "fable": "fable",
                "onyx": "onyx",
                "nova": "nova",
                "shimmer": "shimmer",
            }
    elif request.app.state.config.TTS_ENGINE == "elevenlabs":
        try:
            available_voices = get_elevenlabs_voices(
                api_key=request.app.state.config.TTS_API_KEY
            )
        except Exception:
            # Avoided @lru_cache with exception
            pass
    elif request.app.state.config.TTS_ENGINE == "azure":
        try:
            region = request.app.state.config.TTS_AZURE_SPEECH_REGION
            base_url = request.app.state.config.TTS_AZURE_SPEECH_BASE_URL
            url = (
                base_url or f"https://{region}.tts.speech.microsoft.com"
            ) + "/cognitiveservices/voices/list"
            headers = {
                "Ocp-Apim-Subscription-Key": request.app.state.config.TTS_API_KEY
            }

            response = requests.get(url, headers=headers)
            response.raise_for_status()
            voices = response.json()

            for voice in voices:
                available_voices[voice["ShortName"]] = (
                    f"{voice['DisplayName']} ({voice['ShortName']})"
                )
        except requests.RequestException as e:
            log.error(f"Error fetching voices: {str(e)}")

    return available_voices


@lru_cache
def get_elevenlabs_voices(api_key: str) -> dict:
    """
    Note, set the following in your .env file to use Elevenlabs:
    AUDIO_TTS_ENGINE=elevenlabs
    AUDIO_TTS_API_KEY=sk_...  # Your Elevenlabs API key
    AUDIO_TTS_VOICE=EXAVITQu4vr4xnSDxMaL  # From https://api.elevenlabs.io/v1/voices
    AUDIO_TTS_MODEL=eleven_multilingual_v2
    """

    try:
        # TODO: Add retries
        response = requests.get(
            "https://api.elevenlabs.io/v1/voices",
            headers={
                "xi-api-key": api_key,
                "Content-Type": "application/json",
            },
        )
        response.raise_for_status()
        voices_data = response.json()

        voices = {}
        for voice in voices_data.get("voices", []):
            voices[voice["voice_id"]] = voice["name"]
    except requests.RequestException as e:
        # Avoid @lru_cache with exception
        log.error(f"Error fetching voices: {str(e)}")
        raise RuntimeError(f"Error fetching voices: {str(e)}")

    return voices


@router.get("/voices")
async def get_voices(request: Request, user=Depends(get_verified_user)):
    return {
        "voices": [
            {"id": k, "name": v} for k, v in get_available_voices(request).items()
        ]
    }

def merge_short_segments(segments, min_duration=5.0, max_duration=30.0):
    """짧은 세그먼트를 의미 있는 단위로 병합"""
    merged = []
    current_segment = None

    for segment in segments:
        duration = segment.end - segment.start

        if current_segment is None:
            current_segment = {
                'start': segment.start,
                'end': segment.end,
                'text': segment.text.strip()
            }
        else:
            # 현재 세그먼트와 병합할지 결정
            combined_duration = segment.end - current_segment['start']

            # 병합 조건:
            # 1. 현재 세그먼트가 너무 짧음 (5초 미만)
            # 2. 병합해도 최대 길이를 초과하지 않음 (30초 미만)
            # 3. 세그먼트 간 간격이 3초 미만
            gap = segment.start - current_segment['end']

            if (duration < min_duration or
                combined_duration < max_duration and gap < 3.0):
                # 세그먼트 병합
                merged_text = current_segment['text'] + " " + segment.text.strip()
                current_segment = {
                    'start': current_segment['start'],
                    'end': segment.end,
                    'text': merged_text
                }
            else:
                # 현재 세그먼트 저장하고 새로운 세그먼트 시작
                merged.append(current_segment)
                current_segment = {
                    'start': segment.start,
                    'end': segment.end,
                    'text': segment.text.strip()
                }

    # 마지막 세그먼트 추가
    if current_segment is not None:
        merged.append(current_segment)

    log.info(f"세그먼트 병합: {len(list(segments))}개 → {len(merged)}개")
    return merged


def filter_repetitive_text(text, max_repeat=2):
    """반복 패턴 및 의미없는 텍스트 제거"""
    if not text or not text.strip():
        return ""

    # 1. 과도한 문자 반복 제거 ("네네네네" → "네")
    text = re.sub(r'(.)\1{3,}', r'\1', text)

    # 2. 단어 반복 제거 ("그니까 그니까 그니까" → "그니까")
    words = text.split()
    filtered_words = []

    for word in words:
        # 연속된 같은 단어 제거
        if len(filtered_words) >= max_repeat:
            recent_words = filtered_words[-max_repeat:]
            if all(w == word for w in recent_words):
                continue
        filtered_words.append(word)

    # 3. 의미없는 추임새나 잡음 제거
    meaningless_patterns = [
        r'\b(응+|어+|음+|그+|뭐+)\b',  # "응응응", "어어어" 등
        r'\b(네+|예+)\s*\b(?=\s*\b(네+|예+)\b)',  # 연속된 "네네네"
        r'\b\w\b(?=\s*\b\w\b\s*\b\w\b)',  # 연속된 한글자 단어 3개 이상
    ]

    result_text = " ".join(filtered_words)
    for pattern in meaningless_patterns:
        result_text = re.sub(pattern, '', result_text)

    # 4. 공백 정리
    result_text = re.sub(r'\s+', ' ', result_text).strip()

    return result_text


def create_enhanced_meeting_document(segments_list):
    """향상된 회의록 문서 생성"""
    doc = Document()

    # 제목 스타일 설정
    title = doc.add_heading('회의록', 0)
    title_format = title.runs[0].font
    title_format.name = '맑은 고딕'
    title_format.size = Pt(18)
    title_format.color.rgb = RGBColor(0, 0, 0)

    # 생성 정보 추가
    info_para = doc.add_paragraph()
    info_para.add_run(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
    info_para.add_run(f"\n총 발언 구간: {len(segments_list)}개")

    # 구분선
    doc.add_paragraph("=" * 60)

    # 세그먼트별 내용 추가
    for i, segment in enumerate(segments_list, 1):
        p = doc.add_paragraph()

        # 시간 정보 (굵게, 파란색)
        time_run = p.add_run(f"[{segment['start']} - {segment['end']}] ")
        time_run.bold = True
        time_run.font.color.rgb = RGBColor(0, 100, 200)
        time_run.font.size = Pt(10)

        # 내용 (일반 텍스트)
        content_run = p.add_run(segment['text'])
        content_run.font.name = '맑은 고딕'
        content_run.font.size = Pt(11)

        # 10개마다 구분선 추가
        if i % 10 == 0 and i < len(segments_list):
            doc.add_paragraph("-" * 40)

    return doc
